from __future__ import annotations

import argparse
import csv
import json
import plistlib
import re
import shutil
from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def clean_text(value: str) -> str:
    value = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", value or "")
    return re.sub(r"\s+", " ", value).strip()


def strip_duration(value: str) -> str:
    return clean_text(re.sub(r"\b\d{2}:\d{2}\b", "", value))


def safe_name(value: str, max_len: int = 80) -> str:
    value = strip_duration(value)
    value = re.sub(r"[\\/:*?\"<>|#%&{}$!'@+`=]", "_", value)
    value = re.sub(r"\s+", "_", value).strip("._ ")
    return (value or "video")[:max_len].strip("._ ")


def norm(value: str) -> str:
    value = strip_duration(value).lower()
    value = re.sub(r"\.(mp4|mov|m4v)\b", "", value)
    value = re.sub(r"[\[\]【】（）()《》〈〉：:，,。.\-—_+/\s]", "", value)
    return value


def add_hyperlink(paragraph, text: str, target: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    r_pr.append(size)

    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def make_webloc(path: Path, url: str) -> None:
    data = {"URL": url}
    with path.open("wb") as f:
        plistlib.dump(data, f)


def prepare_video_records(items: list[dict]) -> list[dict]:
    seen = set()
    records = []
    for item in items:
        title = strip_duration(item.get("cellText", ""))
        url = item.get("src", "")
        if not title or not url:
            continue
        key = (norm(title), url.split("?")[0])
        if key in seen:
            continue
        seen.add(key)
        idx = len(records) + 1
        duration = item.get("duration", "")
        filename = f"{idx:03d}_{safe_name(title)}.webloc"
        records.append(
            {
                "index": idx,
                "title": title,
                "duration": duration,
                "url": url,
                "filename": filename,
                "norm": norm(title),
            }
        )
    return records


def match_record(cell_text: str, records: list[dict]) -> dict | None:
    ncell = norm(cell_text)
    if not ncell:
        return None
    best = None
    best_score = 0
    for rec in records:
        ntitle = rec["norm"]
        if not ntitle:
            continue
        score = 0
        if ntitle in ncell:
            score = len(ntitle)
        elif ncell in ntitle and len(ncell) > 5:
            score = len(ncell)
        else:
            # Useful for cells that include bracketed mp4 filenames or shorter labels.
            chunks = [c for c in re.split(r"\d{2,}|方案|修改|编导版", ntitle) if len(c) >= 3]
            score = sum(len(c) for c in chunks if c in ncell)
        if score > best_score:
            best = rec
            best_score = score
    return best if best_score >= 5 else None


def add_links_to_docx(input_docx: Path, output_docx: Path, records: list[dict]) -> tuple[int, list[dict]]:
    doc = Document(str(input_docx))
    matched_ids = set()
    link_count = 0

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                rec = match_record(cell_text, records)
                if not rec:
                    continue
                if f"打开视频 {rec['index']:03d}" in cell_text:
                    continue
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(0)
                rel_target = f"../Videos/{quote(rec['filename'])}"
                add_hyperlink(p, f"▶ 打开视频 {rec['index']:03d}", rel_target)
                if rec.get("duration"):
                    run = p.add_run(f"  {rec['duration']}")
                    run.font.size = Pt(8)
                matched_ids.add(rec["index"])
                link_count += 1

    unmatched = [r for r in records if r["index"] not in matched_ids]
    if unmatched:
        doc.add_page_break()
        heading = doc.add_paragraph()
        run = heading.add_run("未匹配视频清单")
        run.bold = True
        run.font.size = Pt(16)
        doc.add_paragraph("以下视频已放入 Videos 文件夹链接清单，但没有自动匹配到表格单元格。")
        table = doc.add_table(rows=1, cols=4)
        headers = ["序号", "视频标题", "时长", "链接文件"]
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text
        for rec in unmatched:
            row = table.add_row().cells
            row[0].text = f"{rec['index']:03d}"
            row[1].text = rec["title"]
            row[2].text = rec.get("duration", "")
            row[3].text = rec["filename"]

    doc.save(str(output_docx))
    return link_count, unmatched


def write_manifest(records: list[dict], videos_dir: Path) -> None:
    with (videos_dir / "视频清单.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["序号", "标题", "时长", "链接文件", "飞书视频地址"])
        for rec in records:
            writer.writerow([f"{rec['index']:03d}", rec["title"], rec.get("duration", ""), rec["filename"], rec["url"]])

    (videos_dir / "视频清单.json").write_text(
        json.dumps(records, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--videos-json", type=Path, required=True)
    parser.add_argument("--input-docx", type=Path, required=True)
    parser.add_argument("--package-dir", type=Path, required=True)
    args = parser.parse_args()

    package_dir = args.package_dir
    word_dir = package_dir / "Word"
    videos_dir = package_dir / "Videos"
    word_dir.mkdir(parents=True, exist_ok=True)
    videos_dir.mkdir(parents=True, exist_ok=True)

    records = prepare_video_records(json.loads(args.videos_json.read_text(encoding="utf-8")))
    for rec in records:
        make_webloc(videos_dir / rec["filename"], rec["url"])
    write_manifest(records, videos_dir)

    output_docx = word_dir / "【新版】上影节高速拍摄脚本 -内部_可读优化版_含视频.docx"
    link_count, unmatched = add_links_to_docx(args.input_docx, output_docx, records)

    readme = package_dir / "说明.txt"
    readme.write_text(
        "交付包结构：\\n"
        "Word/ 放可读优化版 Word，表格内已添加相对链接。\\n"
        "Videos/ 放每条视频的 .webloc 链接文件和视频清单。\\n\\n"
        "注意：飞书视频直链需要登录态，普通下载请求返回 401；当前 Chrome 页面在尝试 UI 下载时崩溃，"
        "因此本次交付包使用可打开的飞书视频链接文件代替真实 MP4。若需要真实 MP4，请保持 Chrome 页面轻量加载后分批手动点击飞书视频卡片下载按钮，"
        "再把下载得到的 MP4 按 视频清单.csv 的序号放入 Videos/。\\n\\n"
        f"采集到视频条目：{len(records)}\\n"
        f"Word 中添加链接：{link_count}\\n"
        f"未匹配条目：{len(unmatched)}\\n",
        encoding="utf-8",
    )

    print(json.dumps({"records": len(records), "link_count": link_count, "unmatched": len(unmatched), "output_docx": str(output_docx)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
